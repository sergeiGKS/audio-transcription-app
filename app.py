import streamlit as st
import tempfile
import whisper
import io
from docx import Document
import os

# Optional: Configure the page
st.set_page_config(page_title="Application de retranscription d'audio en texte", layout="centered")

# App Title and description
st.title("French Audio Transcription App")
st.markdown(
    """
 Téléchargez un fichier audio en français et laissez l'application le transcrire à l'aide du modèle Whisper d'OpenAI. 
 Une fois la transcription terminée, un document Word sera généré et disponible en téléchargement.
    """
)

    

# File uploader widget that accepts common audio file formats
uploaded_files = st.file_uploader("Téléchargez un ou plusieurs fichiers audio", 
                                 type=["wav", "mp3", "m4a", "ogg", "flac", "opus"],
                                 accept_multiple_files=True)

# Store temporary file info as a list of tuples (original filename, temporary file path)
temp_files = []

if uploaded_files:
    for uploaded_file in uploaded_files:
        st.audio(uploaded_file)
        # Save each uploaded file to a temporary file so that Whisper can process it
        with tempfile.NamedTemporaryFile(delete=False, suffix=".wav") as tmp_file:
            tmp_file.write(uploaded_file.read())
            temp_files.append((uploaded_file.name, tmp_file.name))
            st.write(f"{uploaded_file.name}")


if st.button("Transcrire"):
    
    st.info("Transcription en cours...")

    # Load the Whisper model (choose the desired size, e.g., "base", "small", "medium", etc.)
    model = whisper.load_model("base")


    # Create a new Word document
    doc = Document()
    doc.add_heading("Transcription Audio. par sergeghomsi@gmail.com", 0)


    # Process each temporary file and add its transcription as a section
    for file_name, tmp_file_path in temp_files:
        st.info(f"Transcription du fichier : {file_name}")
        result = model.transcribe(tmp_file_path, language="fr")
        transcription = result.get("text", "Aucune transcription disponible.")
        
        st.subheader(f"Transcription de {file_name}")
        st.write(transcription)

        # Add a heading and paragraph for each file in the Word document
        doc.add_heading(f"Transcription de {file_name}", level=1)
        doc.add_paragraph(transcription)


        # Delete the temporary audio file after processing
        if os.path.exists(tmp_file_path):
            os.remove(tmp_file_path)

    # Save the Word document to an in-memory BytesIO stream
    word_buffer = io.BytesIO()
    doc.save(word_buffer)
    word_buffer.seek(0)



    # Provide a download button for the Word document
    st.download_button(
        label="Télécharger le document Word",
        data=word_buffer,
        file_name="transcription.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    


